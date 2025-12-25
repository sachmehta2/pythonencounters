#!/usr/bin/env python3
"""url_crawl.py

Crawl a website (same-domain links) and export extracted text into a Word document.

Safety / guardrails:
  - max pages
  - max depth
  - request delay
  - same-domain only
"""

import argparse
import time
from collections import deque
from urllib.parse import urljoin, urldefrag, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document


DEFAULT_UA = "pythonencounters-url-crawl/1.0 (+local research; respect site terms)"


def normalize_url(url: str) -> str:
    url, _frag = urldefrag(url)
    return url.rstrip("/")


def same_domain(base: str, candidate: str) -> bool:
    b = urlparse(base)
    c = urlparse(candidate)
    return (c.scheme in {"http", "https"}) and (c.netloc == b.netloc)


def extract_text(session: requests.Session, url: str, timeout: int) -> str | None:
    try:
        r = session.get(url, timeout=timeout)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        # Remove scripts/styles
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)
        return text
    except requests.RequestException as e:
        print(f"Error fetching {url}: {e}")
        return None


def get_sublinks(session: requests.Session, base_url: str, url: str, timeout: int) -> set[str]:
    links: set[str] = set()
    try:
        r = session.get(url, timeout=timeout)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for a in soup.find_all("a", href=True):
            full_url = urljoin(base_url, a["href"])
            full_url = normalize_url(full_url)
            if same_domain(base_url, full_url):
                links.add(full_url)
    except requests.RequestException as e:
        print(f"Error fetching sublinks from {url}: {e}")
    return links


def crawl_and_extract(
    start_url: str,
    output_file: str,
    max_pages: int,
    max_depth: int,
    delay_seconds: float,
    timeout: int,
) -> None:
    start_url = normalize_url(start_url)
    doc = Document()
    visited: set[str] = set()

    q = deque([(start_url, 0)])
    session = requests.Session()
    session.headers.update({"User-Agent": DEFAULT_UA})

    pages_written = 0

    while q and pages_written < max_pages:
        url, depth = q.popleft()
        if url in visited:
            continue
        visited.add(url)

        print(f"Crawling (depth {depth}): {url}")
        text = extract_text(session, url, timeout=timeout)
        if text:
            doc.add_heading(url, level=1)
            doc.add_paragraph(text)
            pages_written += 1

        if depth < max_depth:
            sublinks = get_sublinks(session, start_url, url, timeout=timeout)
            for link in sublinks:
                if link not in visited:
                    q.append((link, depth + 1))

        if delay_seconds > 0:
            time.sleep(delay_seconds)

    doc.save(output_file)
    print(f"Done. Pages written: {pages_written}. Output saved to {output_file}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Crawl a site and export page text to a DOCX.")
    parser.add_argument("--start-url", help="Start URL (e.g., https://example.com).")
    parser.add_argument("--output", help="Output .docx filename.")
    parser.add_argument("--max-pages", type=int, default=50, help="Maximum pages to crawl (default: 50).")
    parser.add_argument("--max-depth", type=int, default=2, help="Maximum link depth (default: 2).")
    parser.add_argument("--delay", type=float, default=0.5, help="Delay between requests in seconds (default: 0.5).")
    parser.add_argument("--timeout", type=int, default=10, help="HTTP timeout seconds (default: 10).")
    args = parser.parse_args()

    start_url = args.start_url or input("Enter the base URL to crawl (e.g., https://example.com): ").strip()
    if not start_url.startswith(("http://", "https://")):
        print("Invalid URL. Please include 'http://' or 'https://'.")
        return 1

    output_file = args.output or input("Enter the output Word filename (e.g., output.docx): ").strip()
    if not output_file:
        output_file = "output.docx"
    if not output_file.endswith(".docx"):
        output_file += ".docx"

    crawl_and_extract(
        start_url=start_url,
        output_file=output_file,
        max_pages=max(1, args.max_pages),
        max_depth=max(0, args.max_depth),
        delay_seconds=max(0.0, args.delay),
        timeout=max(1, args.timeout),
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
